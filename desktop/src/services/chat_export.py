"""Export ticket chat history to PDF and Word (.docx) formats."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any, Callable, Dict, List, Optional
import html
import re

AttachmentFetcher = Optional[Callable[[int], Optional[bytes]]]

REPLY_MARKERS = (
    "-----original message-----",
    "----- forwarded message -----",
    "________________________________",
    "from:",
    "sent from my iphone",
    "sent from my android",
)


def _strip_email_quotes(text: str) -> str:
    if not text:
        return ""

    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    cleaned: List[str] = []

    for line in lines:
        stripped = line.strip()
        lower = stripped.lower()

        if stripped.startswith(">"):
            break
        if " wrote:" in lower and lower.startswith("on "):
            break
        if any(marker in lower for marker in REPLY_MARKERS):
            break

        cleaned.append(line)

    return "\n".join(cleaned).strip()


def _format_export_timestamp(timestamp: str) -> str:
    if not timestamp:
        return ""
    try:
        normalized = timestamp.replace("Z", "+00:00")
        dt = datetime.fromisoformat(normalized)
        return dt.strftime("%d %b %Y at %I:%M %p").lstrip("0")
    except ValueError:
        return timestamp.replace("T", " ")[:19]


def _author_name(comment: Dict[str, Any]) -> str:
    author = comment.get("author") or {}
    return author.get("full_name") or author.get("username") or "Unknown"


def default_export_filename(task_data: Dict[str, Any], extension: str) -> str:
    task_id = task_data.get("id", "ticket")
    title = task_data.get("title") or "untitled"
    safe = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "_")
    safe = safe[:40] or "ticket"
    return f"PM-{task_id}_{safe}_chat.{extension}"


def _assignee_text(task_data: Dict[str, Any]) -> str:
    assignees = task_data.get("assignees") or []
    if not assignees:
        return "None"
    parts = []
    for user in assignees:
        name = user.get("full_name") or user.get("username") or "User"
        email = user.get("email")
        parts.append(f"{name} ({email})" if email else name)
    return ", ".join(parts)


def _due_date_text(task_data: Dict[str, Any]) -> str:
    due_date = task_data.get("due_date")
    if not due_date:
        return "None"
    return _format_export_timestamp(str(due_date))


def _task_summary_lines(task_data: Dict[str, Any]) -> List[str]:
    return [
        f"Ticket: PM-{task_data.get('id', '')}",
        f"Title: {task_data.get('title', 'Untitled')}",
        f"Status: {(task_data.get('status') or 'todo').replace('_', ' ').title()}",
        f"Priority: {(task_data.get('priority') or 'medium').title()}",
        f"Due date: {_due_date_text(task_data)}",
        f"Assignees: {_assignee_text(task_data)}",
        "",
        "Description:",
        (task_data.get("description") or "(No description)").strip(),
        "",
        f"Chat history ({task_data.get('_comment_count', 0)} messages):",
        "",
    ]


def _prepare_comments(comments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return sorted(comments or [], key=lambda item: item.get("created_at") or "")


def export_chat_history_pdf(
    path: str,
    task_data: Dict[str, Any],
    comments: List[Dict[str, Any]],
    attachment_fetcher: AttachmentFetcher = None,
) -> None:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    ordered = _prepare_comments(comments)
    task_data = {**task_data, "_comment_count": len(ordered)}

    def escape(text: str) -> str:
        return html.escape(text or "").replace("\n", "<br/>")

    doc = SimpleDocTemplate(
        path,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=f"PM-{task_data.get('id')} chat history",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TicketTitle",
        parent=styles["Title"],
        alignment=TA_LEFT,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "TicketMeta",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    comment_meta_style = ParagraphStyle(
        "CommentMeta",
        parent=styles["Heading3"],
        fontSize=11,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "CommentBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    attachment_style = ParagraphStyle(
        "AttachmentNote",
        parent=styles["Italic"],
        fontSize=9,
        textColor="#555555",
        spaceAfter=4,
    )

    story = []
    story.append(
        Paragraph(
            f"PM-{task_data.get('id')}: {escape(task_data.get('title', 'Untitled'))}",
            title_style,
        )
    )
    for line in _task_summary_lines(task_data):
        if line:
            story.append(Paragraph(escape(line), meta_style))
        else:
            story.append(Spacer(1, 6))

    if not ordered:
        story.append(Paragraph("No comments on this ticket.", body_style))
    else:
        for index, comment in enumerate(ordered, start=1):
            source = " (via email)" if comment.get("source") == "email" else ""
            story.append(
                Paragraph(
                    (
                        f"{index}. <b>{escape(_author_name(comment))}</b>"
                        f" — {escape(_format_export_timestamp(comment.get('created_at', '')))}{source}"
                    ),
                    comment_meta_style,
                )
            )
            body = _strip_email_quotes(comment.get("body") or "")
            attachments = comment.get("attachments") or []

            if body:
                story.append(Paragraph(escape(body), body_style))
            elif attachments:
                story.append(Paragraph("<i>Shared an image</i>", body_style))

            for attachment in attachments:
                filename = attachment.get("filename") or "attachment"
                story.append(Paragraph(f"Attachment: {escape(filename)}", attachment_style))
                attachment_id = attachment.get("id")
                content_type = attachment.get("content_type") or ""
                if (
                    attachment_fetcher
                    and attachment_id
                    and content_type.startswith("image/")
                ):
                    data = attachment_fetcher(attachment_id)
                    if data:
                        try:
                            story.append(
                                RLImage(
                                    BytesIO(data),
                                    width=4 * inch,
                                    height=3 * inch,
                                    kind="proportional",
                                )
                            )
                        except Exception:
                            pass
            story.append(Spacer(1, 8))

    doc.build(story)


def export_chat_history_docx(
    path: str,
    task_data: Dict[str, Any],
    comments: List[Dict[str, Any]],
    attachment_fetcher: AttachmentFetcher = None,
) -> None:
    from docx import Document
    from docx.shared import Inches

    ordered = _prepare_comments(comments)
    task_data = {**task_data, "_comment_count": len(ordered)}

    document = Document()
    document.add_heading(
        f"PM-{task_data.get('id')}: {task_data.get('title', 'Untitled')}",
        level=0,
    )

    for line in _task_summary_lines(task_data):
        if line.startswith("Chat history"):
            document.add_heading(line, level=1)
        elif line in ("Description:", ""):
            if line:
                document.add_paragraph(line)
        else:
            document.add_paragraph(line)

    if not ordered:
        document.add_paragraph("No comments on this ticket.")
    else:
        for index, comment in enumerate(ordered, start=1):
            source = " (via email)" if comment.get("source") == "email" else ""
            document.add_heading(
                f"{index}. {_author_name(comment)}",
                level=2,
            )
            meta = document.add_paragraph()
            meta.add_run(
                f"{_format_export_timestamp(comment.get('created_at', ''))}{source}"
            ).italic = True

            body = _strip_email_quotes(comment.get("body") or "")
            attachments = comment.get("attachments") or []

            if body:
                document.add_paragraph(body)
            elif attachments:
                document.add_paragraph("Shared an image").italic = True

            for attachment in attachments:
                filename = attachment.get("filename") or "attachment"
                document.add_paragraph(f"Attachment: {filename}").italic = True
                attachment_id = attachment.get("id")
                content_type = attachment.get("content_type") or ""
                if (
                    attachment_fetcher
                    and attachment_id
                    and content_type.startswith("image/")
                ):
                    data = attachment_fetcher(attachment_id)
                    if data:
                        try:
                            document.add_picture(BytesIO(data), width=Inches(4.5))
                        except Exception:
                            pass

    document.save(path)
